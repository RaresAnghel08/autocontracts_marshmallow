from flask import Flask, render_template, request, send_file
from docx import Document
from io import BytesIO
from docx.shared import Inches
import base64
from PIL import Image
from docx2pdf import convert
import shutil
import subprocess
import io, datetime, os, re, copy, zipfile
import logging
try:
    import pythoncom
except Exception:
    pythoncom = None

app = Flask(__name__)


# Note: HTML moved to templates/form.html and static/styles.css

@app.route('/')
def form():
    return render_template('form.html')

@app.route('/generate', methods=['POST'])
def generate_docx():
    try:
        logging.info("Generate request received")
        data = request.form.to_dict()
        logging.info(f"Form data keys: {list(data.keys())}")

        # Generare automată număr și dată contract
        today = datetime.datetime.now()
        data['data_contract'] = today.strftime('%d.%m.%Y')
        data['numar_contract'] = f"MARSHMALLOW-{today.strftime('%Y%m%d-%H%M%S')}"

        # Creăm un folder temporar sigur în proiect per copil
        base_temp = os.path.join(os.getcwd(), 'temp')
        os.makedirs(base_temp, exist_ok=True)

        # Construim numele folderului din nume + prenume copil și îl sanitizăm
        child_raw = f"{data.get('nume_copil','').strip()}_{data.get('prenume_copil','').strip()}".strip('_ ')
        def _sanitize(s):
            s = s.strip()
            s = re.sub(r'[^A-Za-z0-9_-]', '_', s)
            s = re.sub(r'_+', '_', s)
            return s or 'copil'

        child_name = _sanitize(child_raw)

        # Dacă folderul există deja, adăugăm sufix _1, _2, ...
        candidate = os.path.join(base_temp, child_name)
        count = 1
        while os.path.exists(candidate):
            candidate = os.path.join(base_temp, f"{child_name}_{count}")
            count += 1

        temp_dir = candidate
        os.makedirs(temp_dir, exist_ok=True)

        # Move any stray files from the base temp folder into this child's folder
        try:
            for entry in os.listdir(base_temp):
                entry_path = os.path.join(base_temp, entry)
                # skip directories (including other child folders)
                if os.path.isfile(entry_path):
                    # move files (e.g., leftover .docx/.pdf/.png) into the child folder
                    shutil.move(entry_path, os.path.join(temp_dir, entry))
        except Exception:
            # don't fail the whole request if cleanup/move can't run
            pass
        signature_path = os.path.join(temp_dir, 'signature.png')

        # Salvăm semnătura ca imagine
        try:
            signature_parts = data['signature_data'].split(',')
            if len(signature_parts) >= 2:
                signature_data = signature_parts[1]
                signature_bytes = base64.b64decode(signature_data)
                signature_img = Image.open(io.BytesIO(signature_bytes))
                signature_img.save(signature_path)
            else:
                # No valid signature data, create empty file or skip
                with open(signature_path, 'wb') as f:
                    f.write(b'')
        except Exception as e:
            logging.warning(f"Failed to process signature: {e}")
            # Create empty signature file
            with open(signature_path, 'wb') as f:
                f.write(b'')

        # Template-uri DOCX (use project-relative paths)
        base_dir = os.path.dirname(os.path.abspath(__file__))
        templates_dir = os.path.join(base_dir, 'templates')
        contracts = [
            (os.path.join(templates_dir, 'educational.docx'), "contract_educational_completat.docx"),
            (os.path.join(templates_dir, 'catering.docx'), "contract_catering_completat.docx")
        ]

        logging.info(f"Templates directory: {templates_dir}")
        logging.info(f"Template files exist: educational={os.path.exists(os.path.join(templates_dir, 'educational.docx'))}, catering={os.path.exists(os.path.join(templates_dir, 'catering.docx'))}")



        generated_files = []
        failed_conversions = []
        
        # compile a regex to catch {{semnatura}} with optional spaces and case-insensitive
        sig_pattern = re.compile(r"\{\{\s*semnatura\s*\}\}", flags=re.IGNORECASE)
        for template_path, output_name in contracts:
            doc = Document(template_path)
            signature_inserted = False

            def replace_placeholders_in_paragraph(p):
                nonlocal signature_inserted
                # replace normal placeholders from form data
                # skip numeric placeholders handled specially (e.g. '3', '4', '5')
                for key, val in data.items():
                    if key not in ('signature_data', '3', '4', '5'):
                        if f'{{{{{key}}}}}' in p.text:
                            p.text = p.text.replace(f'{{{{{key}}}}}', val)

                # special for {{4}} program selection: use checkbox characters so boxes remain visible
                if '{{4}}' in p.text:
                    prog = data.get('program', '')
                    # use checked/unchecked box unicode characters
                    checked = '☑'
                    unchecked = '☐'
                    if prog == 'normal':
                        # first {{4}} -> checked, second -> unchecked
                        p.text = p.text.replace('{{4}}', checked, 1).replace('{{4}}', unchecked, 1)
                    elif prog == 'prelungit':
                        # first -> unchecked, second -> checked
                        p.text = p.text.replace('{{4}}', unchecked, 1).replace('{{4}}', checked, 1)
                    else:
                        p.text = p.text.replace('{{4}}', unchecked)

                # special for {{5}} group selection: render checked/unchecked for four group options
                if '{{5}}' in p.text:
                    grp = data.get('5', '').lower()
                    checked = '☑'
                    unchecked = '☐'
                    # expected four placeholders in template for the four groups
                    if p.text.count('{{5}}') >= 4:
                        order = ['mica', 'mica_b', 'mijlocie', 'mare']
                        new_text = p.text
                        for opt in order:
                            if grp == opt:
                                new_text = new_text.replace('{{5}}', checked, 1)
                            else:
                                new_text = new_text.replace('{{5}}', unchecked, 1)
                        p.text = new_text
                    else:
                        # fallback: single placeholder -> checked if any group chosen
                        p.text = p.text.replace('{{5}}', checked if grp else unchecked)

                # special for {{3}} consent: use checkbox characters and preserve labels
                if '{{3}}' in p.text:
                    consent = data.get('3', '').lower()
                    checked = '☑'
                    unchecked = '☐'
                    is_agree = consent in ('da', 'sunt', 'sunt de acord', 'true', 'on', 'yes')
                    # If template contains two placeholders (one per option), mark first/second appropriately
                    if p.text.count('{{3}}') >= 2:
                        if is_agree:
                            # first -> checked, second -> unchecked
                            p.text = p.text.replace('{{3}}', checked, 1).replace('{{3}}', unchecked, 1)
                        else:
                            # first -> unchecked, second -> checked
                            p.text = p.text.replace('{{3}}', unchecked, 1).replace('{{3}}', checked, 1)
                    else:
                        # single placeholder: replace with checked/unchecked
                        p.text = p.text.replace('{{3}}', checked if is_agree else unchecked)

                # check for signature placeholder variants
                if sig_pattern.search(p.text):
                    # remove placeholder
                    p.text = sig_pattern.sub('', p.text)
                    try:
                        # add picture which creates a new paragraph with the image
                        if os.path.exists(signature_path) and os.path.getsize(signature_path) > 0:
                            width = Inches(2)
                            doc.add_picture(signature_path, width=width)
                            pic_para = doc.paragraphs[-1]
                            # get the drawing element from the picture run
                            pic_run = pic_para.runs[0] if pic_para.runs else None
                            drawing_elem = None
                            if pic_run is not None:
                                for child in pic_run._r:
                                    # look for drawing element (namespace-aware)
                                    if 'drawing' in child.tag:
                                        drawing_elem = child
                                        break
                            if drawing_elem is not None:
                                # create a new run in the target paragraph and insert the drawing
                                new_run = p.add_run()
                                new_run._r.append(copy.deepcopy(drawing_elem))
                                # move the new run to the start of the paragraph so image appears before text
                                p._p.insert(0, new_run._r)
                                # remove the temporary picture paragraph
                                try:
                                    pic_para._p.getparent().remove(pic_para._p)
                                except Exception:
                                    pass
                                signature_inserted = True
                    except Exception as e:
                        logging.warning(f"Failed to insert signature image: {e}")

            # process top-level paragraphs
            for p in doc.paragraphs:
                replace_placeholders_in_paragraph(p)

            # process paragraphs inside table cells as well
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_placeholders_in_paragraph(p)

            # process textboxes in shapes
            for shape in doc.inline_shapes:
                if hasattr(shape, 'has_text_frame') and shape.has_text_frame:
                    for p in shape.text_frame.paragraphs:
                        replace_placeholders_in_paragraph(p)

            # If no {{semnatura}} was present/inserted, add signature block at end
            if not signature_inserted:
                doc.add_paragraph(f"\nSemnătură părinte:\n{data.get('nume_mama', '')} / {data.get('nume_tata', '')}")
                try:
                    if os.path.exists(signature_path) and os.path.getsize(signature_path) > 0:
                        width = Inches(2)
                        doc.add_picture(signature_path, width=width)
                except Exception as e:
                    logging.warning(f"Failed to add signature at end: {e}")
            output_path = os.path.join(temp_dir, output_name)
            doc.save(output_path)

            # Try to convert DOCX to PDF
            pdf_name = os.path.splitext(output_name)[0] + '.pdf'
            pdf_path = os.path.join(temp_dir, pdf_name)

            def try_convert_to_pdf(input_docx, output_pdf):
                # Try docx2pdf (MS Word COM on Windows)
                try:
                    if pythoncom is not None:
                        try:
                            pythoncom.CoInitialize()
                        except Exception:
                            pass
                        convert(input_docx, output_pdf)
                        return True
                    else:
                        # pythoncom not available (non-Windows); skip docx2pdf
                        logging.info("pythoncom not available; skipping docx2pdf (Word COM)")
                except Exception as e:
                    logging.error(f"docx2pdf failed for {input_docx}: {e}")
                    pass
                # Try LibreOffice/soffice headless conversion
                soffice = shutil.which('soffice') or shutil.which('libreoffice')
                if soffice:
                    try:
                        outdir = os.path.dirname(output_pdf)
                        env = os.environ.copy()
                        env['SAL_USE_JAVA'] = '0'  # Disable Java to avoid javaldx warnings
                        subprocess.run([soffice, '--headless', '--nologo', '--norestore', '--convert-to', 'pdf', '--outdir', outdir, input_docx], check=True, env=env, timeout=30)
                        return os.path.exists(output_pdf)
                    except Exception as e:
                        logging.error(f"soffice failed for {input_docx}: {e}")
                        return False
                return False

            converted = try_convert_to_pdf(output_path, pdf_path)
            if converted:
                file_path = pdf_path
                file_name = pdf_name
                file_type = "PDF"
            else:
                # conversion failed; use DOCX as fallback
                logging.info(f"PDF conversion failed for {output_name}, using DOCX instead")
                failed_conversions.append(output_name)
                file_path = output_path
                file_name = output_name
                file_type = "DOCX"
            
            # Generate a unique ID for this file download
            file_id = f"{child_name}_{os.path.splitext(file_name)[0]}_{int(datetime.datetime.now().timestamp())}"
            generated_files.append({
                'id': file_id,
                'name': file_name,
                'path': file_path,
                'type': file_type
            })

        # Store file info in a global dict (in production, use a database or cache with expiration)
        if not hasattr(app, 'file_downloads'):
            app.file_downloads = {}
        
        for file_info in generated_files:
            app.file_downloads[file_info['id']] = file_info

        # Prepare data for template
        files_for_template = []
        for file_info in generated_files:
            files_for_template.append((file_info['name'], f"/download/{file_info['id']}"))

        # Don't clean up temp directory immediately - files will be served individually
        # Add cleanup logic here if needed for production
        
        return render_template('download.html', files=files_for_template, failed=failed_conversions if failed_conversions else None)
    except Exception as e:
        logging.error(f"Error in generate_docx: {e}")
        return "Eroare la generarea contractelor. Te rugăm să încerci din nou.", 500

@app.route('/download/<file_id>')
def download_file(file_id):
    if not hasattr(app, 'file_downloads'):
        return "File not found", 404
    
    file_info = app.file_downloads.get(file_id)
    if not file_info or not os.path.exists(file_info['path']):
        return "File not found", 404
    
    # Clean up old files (older than 1 hour)
    current_time = datetime.datetime.now().timestamp()
    to_remove = []
    for fid, finfo in app.file_downloads.items():
        # Extract timestamp from file_id (it's at the end)
        try:
            timestamp = float(fid.split('_')[-1])
            if current_time - timestamp > 3600:  # 1 hour
                try:
                    os.remove(finfo['path'])
                except:
                    pass
                to_remove.append(fid)
        except:
            pass
    
    for fid in to_remove:
        del app.file_downloads[fid]
    
    # Determine MIME type
    if file_info['type'] == 'PDF':
        mimetype = 'application/pdf'
    else:
        mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    
    return send_file(file_info['path'], as_attachment=True, download_name=file_info['name'], mimetype=mimetype)

if __name__ == '__main__':
    app.run(debug=True)